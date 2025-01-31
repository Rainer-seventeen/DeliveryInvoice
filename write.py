# -*- coding: gbk -*-
"""������еĹ���json�ļ���ȡ��д��ĺ���"""
import json
import os
from docx import Document

from info import CURRENT_FOLDER
from info import FILE_ENCODING
from operation import Operation
from ChineseUppercase import float_to_chinese_uppercase
from tkinter import messagebox



class Write:
    def __init__(self):
        self.invoice_path = rf"{CURRENT_FOLDER}{r"\invoice.json"}"
        self.framework_path = rf"{CURRENT_FOLDER}{r"\framework.docx"}"
        self.output_path = rf"{CURRENT_FOLDER}{r"\���ɺ��ͻ���.docx"}"

        #��ʼ��invoice.json�ļ�
        if not os.path.exists(self.invoice_path):
            with open(self.invoice_path, 'w', encoding= FILE_ENCODING) as f:
                json.dump({"DeliveryOrders": []}, f, ensure_ascii=False, indent=4)

    def add_new_invoice(self,company_name, company_address, company_phone, company_connector, entry_matrix):
        """
        д��json����ڣ������װ����
        :param company_name: ��˾����
        :param company_address: ��˾��ַ
        :param company_phone: ��˾�绰
        :param company_connector:��˾��ϵ��
        :param entry_matrix:�������ͨ��uiֱ�Ӵ��룩
        """
        op = Operation()
        order_list = op.create_new_order(company_name, company_address, company_phone, company_connector, entry_matrix)
        self.write_invoice_json(order_list)

    def write_invoice_json(self, new_order):
        """
        �����ɵ��ֵ�д��json�ļ�
        :param new_order: �ֵ����͵����ݣ���create_new_order()����
        """
        # ����ļ������ڣ������ļ�����ʼ��Ϊһ���յ� delivery orders ����


        # ��ȡ���е� JSON ����
        with open(self.invoice_path, 'r', encoding= FILE_ENCODING) as f:
            data = json.load(f)

        # ����µ��ͻ�������
        data['DeliveryOrders'].append(new_order)

        # д�� JSON �ļ�
        with open(self.invoice_path, 'w', encoding= FILE_ENCODING) as f:
            json.dump(data, f, ensure_ascii=False, indent=4)

        print("�µ��ͻ�������ӳɹ���")

    def write_docx(self, input_dict):
        """

        :param input_dict: ������ֵ����͵�����
        :return:
        """
        # Step 1: ��ȡ���ݣ����ֵ���
        titled_price = float_to_chinese_uppercase(input_dict["all_total_price"])
        # Step 2: ��ģ���ĵ�
        doc = Document(self.framework_path)

        # Step 3: �滻��ͨ��˾��Ϣ�ֶ�
        # ��������е����е�Ԫ������滻
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if "<id>" in run.text:  # ����Ƿ���ռλ��
                    run.text = run.text.replace("<id>", input_dict["order_no"])
        for table in doc.tables:  # ����ÿ�����
            for row in table.rows:  # ����ÿ��
                for cell in row.cells:  # ����ÿ����Ԫ��
                    # ������Ԫ���ڵĶ�������滻
                    for paragraph in cell.paragraphs:
                        paragraph.text = paragraph.text.replace("<company_name>", input_dict["company_name"])
                        paragraph.text = paragraph.text.replace("<company_address>", input_dict["company_address"])
                        paragraph.text = paragraph.text.replace("<company_phone>", input_dict["company_phone"])
                        paragraph.text = paragraph.text.replace("<connector>", input_dict["company_connector"])
                        paragraph.text = paragraph.text.replace("<date>", input_dict["created_date"])
                        paragraph.text = paragraph.text.replace("<id>", input_dict["order_no"])
                        paragraph.text = paragraph.text.replace("<all_total_price>", str(input_dict["all_total_price"]))
                        paragraph.text = paragraph.text.replace("<titled_total_price>", titled_price)


        # Step 4: �滻���� <xx> ռλ��
        # # �滻�����е�ռλ��
        # for paragraph in doc.paragraphs:
        #     for run in paragraph.runs:
        #         if "<" in run.text:  # ����Ƿ���ռλ��
        #             run.text = replace_product_placeholder(run.text, products)
        #
        # �滻����е�ռλ��
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:  # ������Ԫ���еĶ���
                        for run in paragraph.runs:  # ���������е� Run ����
                            if "<" in run.text:  # ����Ƿ���ռλ��
                                run.text = self.replace_product_placeholder(run.text, input_dict["Products"])

        # Step 5: �����滻����ĵ�
        doc.save(self.output_path)
        print(f"���ɵ����ĵ��ѱ�����: {self.output_path}")
        if messagebox.askyesno('���','�ļ��Ѿ��ɹ����ɣ��Ƿ����ڴ򿪣�'):
            os.startfile(self.output_path)

    @staticmethod
    def replace_product_placeholder(text, input_list):

        products_matrix = []
        for product in input_list:
            product_list = [product["product_name"],product["product_standard"],product["product_unit"],
                    product["quantity"],product["unit_price"],product["product_description"],
                    product["total_price"],product["id"]]

            products_matrix.append(product_list)
        # print(products_matrix)

        import re

        # ����ƥ�� <xx> ��ʽ
        pattern = r"<(\d)(\d)>"  # ƥ�� <00> �� <99> ��ռλ��
        matches = re.findall(pattern, text)

        # �滻����ƥ���ռλ��
        for match in matches:
            row, col = int(match[0]), int(match[1])  # ��ȡ�к�������
            if row < len(products_matrix) and col < len(products_matrix[row]):
                text = text.replace(f"<{row}{col}>", str(products_matrix[row][col]))
            else:
                # �����������Χ���滻Ϊ���ַ���
                text = text.replace(f"<{row}{col}>", "")
        return text
if __name__ == '__main__':

    op = Operation()
    out = op.find_order_by_order_no(order_no= 2025010001)
    wt = Write()
    wt.write_docx(out)
